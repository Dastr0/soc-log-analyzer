"""
Reporter — format output analisa dalam 3 level kedalaman.

Level 1 (default)  : Executive Summary — tabel ringkasan insiden, metrik umum.
Level 2 (--detail)  : + detail per insiden: timeline, metrik, reasoning, sample raw event.
Level 3 (--verbose) : + daftar lengkap event per insiden.

Output ke stdout (bisa di-redirect ke file dengan --output).
"""

import shutil
from datetime import datetime
from typing import List, Optional, TextIO

from src.schema import CommonEvent, Incident


# ─── Terminal Helpers ──────────────────────────────────────────────

def _terminal_width() -> int:
    """Deteksi lebar terminal (fallback 80)."""
    try:
        return shutil.get_terminal_size().columns
    except Exception:
        return 80


def _severity_bar(severity: int) -> str:
    """Bar visual untuk severity: 1-4."""
    colors = {1: "·", 2: "▪", 3: "▨", 4: "■"}
    return colors.get(severity, "?")


def _verdict_tag(verdict: str) -> str:
    """Color-coded label untuk verdict."""
    tags = {
        "CONFIRMED": "⚠ CONFIRMED",
        "LIKELY_TRUE": "⚡ LIKELY_TRUE",
        "LIKELY_FP": "◊ LIKELY_FP",
        "FALSE_POSITIVE": "✓ FALSE_POSITIVE",
        "UNCLEAR": "? UNCLEAR",
    }
    return tags.get(verdict, verdict)


def _truncate(s: str, max_len: int = 40) -> str:
    """Potong string + tambahkan '...' jika terlalu panjang."""
    s = str(s)
    if len(s) <= max_len:
        return s
    return s[:max_len - 3] + "..."


# ─── Output ────────────────────────────────────────────────────────

def report(incidents: List[Incident],
           all_events: List[CommonEvent] = None,
           detail: bool = False,
           verbose: bool = False,
           output_file: Optional[str] = None) -> None:
    """
    Output hasil analisa.

    Args:
        incidents: List Incident.
        all_events: Semua event (buat statistik total).
        detail: Jika True, tampilkan detail per insiden.
        verbose: Jika True, tampilkan daftar event per insiden.
        output_file: Path file output (opsional). Auto-detect format from extension:
                 .docx → Word document, .txt/.md → teks, default → stdout teks.
    """
    # Auto-detect format from extension
    format_docx = output_file and output_file.lower().endswith((".docx", ".docs"))

    if format_docx:
        _write_docx(incidents, all_events, detail or verbose, verbose, output_file)
        return

    out: TextIO = open(output_file, "w", encoding="utf-8") if output_file else None

    try:
        _write_summary(incidents, all_events, out)

        if detail or verbose:
            for incident in incidents:
                _write_incident_detail(incident, out)
                if verbose:
                    _write_incident_events(incident, out)

    finally:
        if out:
            out.close()


def _write_summary(incidents: List[Incident], all_events: List[CommonEvent] = None,
                   out: Optional[TextIO] = None) -> None:
    """Level 1: Executive Summary."""

    def p(text: str = ""):
        print(text)
        if out:
            out.write(text + "\n")

    all_events = all_events or []

    # Statistik
    total = len(all_events)
    severity_counts = {1: 0, 2: 0, 3: 0, 4: 0}
    for e in all_events:
        if e.severity in severity_counts:
            severity_counts[e.severity] += 1

    n_confirmed = sum(1 for i in incidents if i.verdict == "CONFIRMED")
    n_likely_true = sum(1 for i in incidents if i.verdict == "LIKELY_TRUE")
    n_likely_fp = sum(1 for i in incidents if i.verdict == "LIKELY_FP")
    n_fp = sum(1 for i in incidents if i.verdict == "FALSE_POSITIVE")
    n_unknown = sum(1 for i in incidents if i.verdict == "UNCLEAR")

    # Sumber
    sources = set(e.source for e in all_events if e.source)
    source_counts: dict = {}
    for e in all_events:
        source_counts[e.source] = source_counts.get(e.source, 0) + 1
    source_str = ", ".join(
        f"{s}({c:,})" for s, c in sorted(source_counts.items())
    )

    # Rentang waktu
    if all_events:
        t_min = min(e.timestamp for e in all_events)
        t_max = max(e.timestamp for e in all_events)
        time_range = f"{t_min.strftime('%Y-%m-%d %H:%M')} → {t_max.strftime('%H:%M')}"
    else:
        time_range = "N/A"

    w = _terminal_width()
    sep = "═" * min(w, 80)

    p(f"\n{sep}")
    p(f"  EXECUTIVE SUMMARY — soc-log-analyzer v0.2.0")
    p(f"{sep}")
    p(f"  Generated     : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p(f"  Total events  : {total:,}  [{source_str}]")
    p(f"  Time range    : {time_range}")
    p(f"  Severity      : "
      f"CRIT:{severity_counts[4]}  HIGH:{severity_counts[3]}  "
      f"MED:{severity_counts[2]}  LOW:{severity_counts[1]}")
    p(f"  Incidents     : {len(incidents)} "
      f"(CONFIRMED:{n_confirmed}  "
      f"LIKELY_TRUE:{n_likely_true}  "
      f"LIKELY_FP:{n_likely_fp}  "
      f"FP:{n_fp}  UNCLEAR:{n_unknown})")
    p(f"{'─' * min(w, 80)}")

    if not incidents:
        p("  ✓ Tidak ada insiden terdeteksi.")
        return

    # Tabel insiden
    for i, inc in enumerate(incidents, 1):
        dur_str = ""
        if inc.start_time and inc.end_time:
            dur = int((inc.end_time - inc.start_time).total_seconds())
            dur_str = f"{dur // 60}m{dur % 60}s" if dur >= 60 else f"{dur}s"

        sev_max = max((e.severity for e in inc.events), default=1)
        sev_labels = {1: "LOW", 2: "MED", 3: "HIGH", 4: "CRIT"}
        sev_label = sev_labels.get(sev_max, "LOW")

        ips = ", ".join(sorted(inc.unique_ips)[:3]) if inc.unique_ips else "—"
        hosts = ", ".join(sorted(set(e.dst_host for e in inc.events if e.dst_host))[:2]) or "—"

        p(f"\n  ┌ #{i} {inc.pattern} · {_verdict_tag(inc.verdict)} "
          f"({sev_label}, conf {inc.confidence:.0f}%)")
        p(f"  │ {inc.event_count:,} event · {dur_str} · src: {ips} → {hosts}")

        # Kalau ada reasoning, tampilin rekomendasi singkat
        if inc.verdict == "CONFIRMED":
            p(f"  └ ⚠ AKSI: Investigasi segera — pola serangan terkonfirmasi.")
        elif inc.verdict == "LIKELY_TRUE":
            p(f"  └ ⚡ Perlu analisa lanjut — indikasi kuat, tapi ada noise.")
        elif "FP" in inc.verdict:
            p(f"  └ ✓ Cenderung FP — verifikasi manual disarankan.")

    p(f"\n{'─' * min(w, 80)}")
    p(f"  ℹ Level 2 (--detail): detail per insiden + reasoning FP.")
    p(f"  ℹ Level 3 (--verbose): semua event per insiden.")


def _write_incident_detail(incident: Incident, out: Optional[TextIO] = None) -> None:
    """Level 2: Detail per insiden."""
    def p(text: str = ""):
        print(text)
        if out:
            out.write(text + "\n")

    w = min(_terminal_width(), 80)

    p(f"\n╔{'═' * (w - 2)}╗")
    p(f"║  INSIDEN #{incident.id}   {incident.pattern}   "
      f"{incident.verdict}   {incident.confidence:.0f}% confidence"
      f"{' ' * max(0, w - 53 - len(incident.pattern) - len(incident.verdict))}║")
    p(f"╚{'═' * (w - 2)}╝")

    # Timeline
    p("\n── TIMELINE ───────────────────────────────────────────────")
    events_sorted = sorted(incident.events, key=lambda e: e.timestamp)
    shown = min(5, len(events_sorted))
    for e in events_sorted[:shown]:
        ts = e.timestamp.strftime("%H:%M:%S")
        action = e.action or "unknown"
        user = e.user or "?"
        src = f"{e.src_ip}:{e.dst_port}" if e.src_ip and e.dst_port else e.src_ip or "?"
        dst = e.dst_host or e.dst_ip or "?"
        p(f"  {ts} · {action:20s} {user:20s} src:{src} → {dst}")

    if len(events_sorted) > shown:
        p(f"  ... ({len(events_sorted) - shown} event lainnya) ...")
    p(f"── ({incident.start_time.strftime('%H:%M:%S') if incident.start_time else '?'} → "
      f"{incident.end_time.strftime('%H:%M:%S') if incident.end_time else '?'}) "
      f"—— {len(events_sorted)} events")

    # Metrik
    p("\n── METRIK ─────────────────────────────────────────────────")
    dur_str = ""
    if incident.start_time and incident.end_time:
        dur = int((incident.end_time - incident.start_time).total_seconds())
        dur_str = f"{dur // 60} menit {dur % 60} detik" if dur >= 60 else f"{dur} detik"

    p(f"  Total event    : {incident.event_count:,}")
    p(f"  Durasi         : {dur_str}")
    p(f"  Sumber         : {', '.join(sorted(incident.sources))}")
    p(f"  IP terkait     : {', '.join(sorted(incident.unique_ips)) if incident.unique_ips else '—'}")
    p(f"  User terkait   : {', '.join(sorted(incident.unique_users)) if incident.unique_users else '—'}")

    ports = set(e.dst_port for e in incident.events if e.dst_port is not None)
    if ports:
        def _port_key(p):
            try: return int(p)
            except (ValueError, TypeError): return 0
        p(f"  Port target    : {', '.join(str(p) for p in sorted(ports, key=_port_key)[:10])}"
          f"{' (+' + str(len(ports) - 10) + ' lainnya)' if len(ports) > 10 else ''}")

    hosts = set(e.dst_host for e in incident.events if e.dst_host)
    if hosts:
        p(f"  Host target    : {', '.join(sorted(hosts))}")

    processes = set(e.process for e in incident.events if e.process)
    if processes:
        p(f"  Process        : {', '.join(list(processes)[:3])}")

    # Sample raw event
    p("\n── SAMPLE RAW EVENT ───────────────────────────────────────")
    sample = incident.events[0]
    if sample.raw_event:
        import json
        raw_json = json.dumps(sample.raw_event, indent=2, default=str)
        for line in raw_json.split("\n")[:12]:
            p(f"  {line}")
        if len(raw_json.split("\n")) > 12:
            p("  ... (truncated)")

    # Reasoning
    p("\n── FALSE POSITIVE REASONING ───────────────────────────────")
    if incident.reasoning:
        for reason in incident.reasoning:
            p(f"  {reason}")
    else:
        p("  (tidak ada reasoning — data tidak cukup)")

    # Rekomendasi
    p("\n── REKOMENDASI ───────────────────────────────────────────")
    if incident.verdict == "CONFIRMED":
        p(f"  ⚠ INSIDEN TERKONFIRMASI — pola: {incident.pattern}")
        if incident.pattern == "BRUTE_FORCE":
            ips_str = ', '.join(
                sorted(ip for ip in incident.unique_ips if ip and not ip.startswith("10."))[:3])
            if ips_str:
                p(f"    1. Blokir IP: {ips_str}")
            p("    2. Reset password user yang terkena")
            p("    3. Cek Windows EventLog (EventID 4625) untuk konfirmasi")
        elif incident.pattern == "SUSPICIOUS_PROC":
            p("    1. Cek apakah process ini legitimate (tiket change request)")
            p("    2. Ambil sample process + hash untuk analisa")
            p("    3. Isolasi host jika tidak ada justifikasi")
    elif incident.verdict == "FALSE_POSITIVE":
        p("  ✓ KEMUNGKINAN FP TINGGI — tidak perlu tindakan")
        p("    (Tapi tetap review berkala — pastikan ada di trusted_hosts)")
    else:
        p(f"  ? {incident.verdict} — perlu analisa manual")


def _write_incident_events(incident: Incident, out: Optional[TextIO] = None) -> None:
    """Level 3: Daftar lengkap event per insiden."""
    def p(text: str = ""):
        print(text)
        if out:
            out.write(text + "\n")

    p(f"\n─── EVENT LIST (INSIDEN #{incident.id}) ─{'─' * 35}")

    events_sorted = sorted(incident.events, key=lambda e: e.timestamp)
    for i, e in enumerate(events_sorted, 1):
        ts = e.timestamp.strftime("%Y-%m-%d %H:%M:%S")
        sev = _severity_bar(e.severity)
        src = e.src_ip or "?"
        dst = e.dst_host or e.dst_ip or "?"
        action = e.action or "?"
        user = e.user or "?"

        p(f"  {i:4d} | {ts} | {sev} | {action:18s} | user:{user:15s}")
        p(f"       | src:{src} → dst:{dst}  "
          f"port:{e.dst_port or '-'} "
          f"source:{e.source} "
          f"result:{e.result or '-'}")

    p(f"\n  Total: {len(events_sorted)} events untuk insiden #{incident.id}")


# ─── DOCX Export ────────────────────────────────────────────────────

def _write_docx(incidents: List[Incident],
                all_events: List[CommonEvent],
                detail: bool, verbose: bool,
                output_file: str) -> None:
    """Generate report .docx (Word document)."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[!] python-docx belum terinstall. Install: pip install python-docx")
        print("[i] Falling back ke text output...")
        _write_docx_fallback(incidents, all_events, detail, verbose, output_file)
        return

    doc = Document()

    # Title
    title = doc.add_heading("SOC Log Analyzer — Executive Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    # Stats
    total = len(all_events or [])
    n_confirmed = sum(1 for i in incidents if i.verdict == "CONFIRMED")
    n_fp = sum(1 for i in incidents if "FP" in i.verdict)

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(f"Total events: {total:,}")
    doc.add_paragraph(f"Incidents: {len(incidents)} "
                      f"(CONFIRMED: {n_confirmed}, FP: {n_fp})")

    if all_events:
        t_min = min(e.timestamp for e in all_events)
        t_max = max(e.timestamp for e in all_events)
        doc.add_paragraph(f"Time range: {t_min.strftime('%Y-%m-%d %H:%M')} → "
                          f"{t_max.strftime('%H:%M')}")

    # Per-incident sections
    for inc in incidents:
        doc.add_heading(f"Incident #{inc.id}: {inc.pattern} — {inc.verdict} "
                        f"({inc.confidence:.0f}%)", level=2)

        # Key metrics table
        table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
        rows_data = [
            ("Events", str(inc.event_count)),
            ("Source IPs", ", ".join(sorted(inc.unique_ips)[:5]) if inc.unique_ips else "—"),
            ("Users", ", ".join(sorted(inc.unique_users)[:5]) if inc.unique_users else "—"),
            ("Sources", ", ".join(sorted(inc.sources)) if inc.sources else "—"),
            ("Pattern", inc.pattern),
        ]
        for i, (k, v) in enumerate(rows_data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v

        # Reasoning
        if inc.reasoning and detail:
            doc.add_heading("FP Reasoning", level=3)
            for r in inc.reasoning:
                doc.add_paragraph(r, style="List Bullet")

        # Sample events
        if verbose:
            doc.add_heading("Events", level=3)
            for e in sorted(inc.events, key=lambda x: x.timestamp)[:50]:
                doc.add_paragraph(
                    f"{e.timestamp.strftime('%H:%M:%S')} | "
                    f"{e.action or '?'} | src:{e.src_ip or '?'} → "
                    f"{e.dst_host or e.dst_ip or '?'}"
                )

        doc.add_paragraph("")  # spacer

    doc.save(output_file)
    print(f"[✓] DOCX saved: {output_file}")


def _write_docx_fallback(incidents, all_events, detail, verbose, output_file):
    """Fallback: tulis text report dengan ekstensi .txt kalau python-docx nggak ada."""
    fallback_path = output_file.rsplit(".", 1)[0] + ".txt"
    with open(fallback_path, "w") as f:
        for inc in incidents:
            f.write(f"Incident #{inc.id}: {inc.pattern} — {inc.verdict} ({inc.confidence:.0f}%)\n")
            f.write(f"  Events: {inc.event_count}, IPs: {', '.join(sorted(inc.unique_ips)[:5])}\n")
            f.write("\n")
    print(f"[✓] Fallback TXT: {fallback_path} (install python-docx for .docx)")
